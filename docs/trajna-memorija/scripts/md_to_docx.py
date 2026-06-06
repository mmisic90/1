#!/usr/bin/env python3
"""
Konvertor Markdown -> DOCX bez pandoc-a.
Pokriva: H1-H6, bullet list, numbered list, code block (fenced),
inline code, bold, italic, link, tabele (pipe sintaksa), horizontal rule.

Upotreba:
  pip install python-docx
  python3 md_to_docx.py <md_dir> [<docx_dir>]

Primer:
  python3 docs/trajna-memorija/scripts/md_to_docx.py docs/trajna-memorija docs/trajna-memorija/docx
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_font(run, name="Calibri", size=11, bold=False, italic=False, mono=False, color=None):
    if mono:
        run.font.name = "Consolas"
    else:
        run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor(*color)


INLINE_PATTERN = re.compile(
    r"(\*\*([^*]+)\*\*)"          # bold
    r"|(\*([^*]+)\*)"             # italic
    r"|(`([^`]+)`)"               # inline code
    r"|(\[([^\]]+)\]\(([^)]+)\))" # link
)


def add_runs(paragraph, text):
    """Parsira inline markdown u runove."""
    pos = 0
    for m in INLINE_PATTERN.finditer(text):
        if m.start() > pos:
            r = paragraph.add_run(text[pos:m.start()])
            set_font(r)
        if m.group(1):  # bold
            r = paragraph.add_run(m.group(2))
            set_font(r, bold=True)
        elif m.group(3):  # italic
            r = paragraph.add_run(m.group(4))
            set_font(r, italic=True)
        elif m.group(5):  # inline code
            r = paragraph.add_run(m.group(6))
            set_font(r, mono=True, size=10, color=(0xC7, 0x25, 0x4E))
        elif m.group(7):  # link
            r = paragraph.add_run(m.group(8))
            set_font(r, color=(0x12, 0x55, 0xCC))
            r.font.underline = True
        pos = m.end()
    if pos < len(text):
        r = paragraph.add_run(text[pos:])
        set_font(r)


def add_heading(doc, text, level):
    p = doc.add_paragraph()
    sizes = {1: 22, 2: 18, 3: 15, 4: 13, 5: 12, 6: 11}
    r = p.add_run(text)
    set_font(r, size=sizes.get(level, 11), bold=True)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)


def add_code_block(doc, lines):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), 'F2F2F2')
    pPr.append(shd)
    r = p.add_run("\n".join(lines))
    set_font(r, mono=True, size=10)


def add_table(doc, rows):
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=n_cols)
    table.style = 'Light Grid Accent 1'
    for i, row in enumerate(rows):
        for j, cell in enumerate(row):
            tc = table.rows[i].cells[j]
            tc.text = ""
            p = tc.paragraphs[0]
            add_runs(p, cell.strip())
            if i == 0:
                for run in p.runs:
                    run.font.bold = True


def parse_table_block(block_lines):
    rows = []
    for ln in block_lines:
        ln = ln.strip()
        if not ln:
            continue
        if re.match(r'^\|?\s*[-:]+\s*(\|\s*[-:]+\s*)+\|?$', ln):
            continue
        ln = ln.strip().strip('|')
        cells = [c.strip() for c in ln.split('|')]
        rows.append(cells)
    return rows


def convert(md_path: Path, docx_path: Path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    text = md_path.read_text(encoding='utf-8')
    lines = text.split('\n')

    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.rstrip()

        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].rstrip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            add_code_block(doc, code_lines)
            i += 1
            continue

        m = re.match(r'^(#{1,6})\s+(.*)$', stripped)
        if m:
            level = len(m.group(1))
            add_heading(doc, m.group(2), level)
            i += 1
            continue

        if re.match(r'^-{3,}\s*$', stripped) or re.match(r'^\*{3,}\s*$', stripped):
            p = doc.add_paragraph()
            r = p.add_run('—' * 40)
            set_font(r, color=(0xAA, 0xAA, 0xAA))
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            i += 1
            continue

        if '|' in stripped and i + 1 < len(lines) and re.match(r'^\|?\s*[-:]+\s*(\|\s*[-:]+\s*)+\|?$', lines[i+1].strip()):
            block = [stripped]
            i += 1
            while i < len(lines) and '|' in lines[i]:
                block.append(lines[i])
                i += 1
            rows = parse_table_block(block)
            add_table(doc, rows)
            continue

        m = re.match(r'^(\s*)[-*+]\s+(.*)$', stripped)
        if m:
            indent = len(m.group(1)) // 2
            p = doc.add_paragraph(style='List Bullet')
            p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
            add_runs(p, m.group(2))
            i += 1
            continue

        m = re.match(r'^(\s*)\d+\.\s+(.*)$', stripped)
        if m:
            indent = len(m.group(1)) // 2
            p = doc.add_paragraph(style='List Number')
            p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
            add_runs(p, m.group(2))
            i += 1
            continue

        m = re.match(r'^>\s?(.*)$', stripped)
        if m:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            r = p.add_run('| ')
            set_font(r, color=(0x88, 0x88, 0x88))
            add_runs(p, m.group(1))
            i += 1
            continue

        if not stripped:
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    doc.save(str(docx_path))


def main():
    if len(sys.argv) < 2:
        print("Upotreba: md_to_docx.py <md_dir> [<docx_dir>]")
        sys.exit(1)
    md_dir = Path(sys.argv[1])
    docx_dir = Path(sys.argv[2]) if len(sys.argv) > 2 else md_dir / 'docx'
    docx_dir.mkdir(parents=True, exist_ok=True)

    md_files = sorted(md_dir.glob('*.md'))
    for md in md_files:
        if md.parent.name == 'scripts':
            continue
        out = docx_dir / (md.stem + '.docx')
        print(f"Konvertujem: {md.name} -> {out.relative_to(md_dir)}")
        convert(md, out)
    print(f"Gotovo: {len(md_files)} fajl(ova).")


if __name__ == '__main__':
    main()
