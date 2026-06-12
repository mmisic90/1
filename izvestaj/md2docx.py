#!/usr/bin/env python3
"""Konvertuje Markdown (.md) u .docx — zajednički generator za sve podneske/analize.
Upotreba: python3 md2docx.py ulaz.md [izlaz.docx]
"""
import re, sys
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BLACK = RGBColor(0, 0, 0)


def blacken(p):
    for r in p.runs:
        r.font.color.rgb = BLACK


def add_runs(p, text):
    parts = re.split(r'(\*\*.+?\*\*|\*.+?\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith('*') and part.endswith('*'):
            p.add_run(part[1:-1]).italic = True
        else:
            p.add_run(part.replace('\\_', '_'))


def convert(src_path, out_path):
    src = open(src_path, encoding='utf-8').read()
    doc = Document()
    st = doc.styles['Normal']
    st.font.name = 'Times New Roman'
    st.font.size = Pt(12)
    for sec in doc.sections:
        sec.left_margin = Cm(2.54)
        sec.right_margin = Cm(2.54)
        sec.top_margin = Cm(2.54)
        sec.bottom_margin = Cm(2.54)

    lines = src.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if line.startswith('|'):
            tbl = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                if not set(''.join(row)) <= set('-: '):
                    tbl.append(row)
                i += 1
            if tbl:
                t = doc.add_table(rows=len(tbl), cols=len(tbl[0]))
                t.style = 'Table Grid'
                for ri, row in enumerate(tbl):
                    for ci, cell in enumerate(row):
                        add_runs(t.cell(ri, ci).paragraphs[0], cell)
            continue
        if line.strip() == '---':
            i += 1
            continue
        if line.startswith('[R] '):
            p = doc.add_paragraph()
            add_runs(p, line[4:])
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            i += 1
            continue
        if line.startswith('> '):
            p = doc.add_paragraph()
            add_runs(p, line[2:])
            p.paragraph_format.left_indent = Cm(1.0)
            p.style = doc.styles['Quote'] if 'Quote' in [s.name for s in doc.styles] else p.style
        elif line.startswith('### '):
            p = doc.add_heading('', level=3)
            add_runs(p, line[4:])
            blacken(p)
        elif line.startswith('## '):
            p = doc.add_heading('', level=2)
            add_runs(p, line[3:])
            blacken(p)
        elif line.startswith('# '):
            p = doc.add_heading('', level=1)
            add_runs(p, line[2:])
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            blacken(p)
        elif re.match(r'^\d+\.\s', line):
            p = doc.add_paragraph()
            add_runs(p, line)
            p.paragraph_format.left_indent = Cm(0.75)
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, line[2:])
        elif line.strip():
            p = doc.add_paragraph()
            add_runs(p, line)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        i += 1

    doc.save(out_path)
    print('saved', out_path)


if __name__ == '__main__':
    src = sys.argv[1]
    out = sys.argv[2] if len(sys.argv) > 2 else src.rsplit('.', 1)[0] + '.docx'
    convert(src, out)
